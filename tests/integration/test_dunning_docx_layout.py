"""Layout-Invarianten des Mahn-DOCX.

Das Word-Dokument brach auf zwei Seiten um, obwohl der Inhalt auf eine passt.
Zwei Ursachen, beide hier abgesichert:

1. **Gestapelte Bloecke.** Infoblock (Rechnungsnr./Kundennr./Daten) und
   Empfaengeranschrift standen als volle Zeilen untereinander statt
   nebeneinander. Word kennt kein ``float`` wie das PDF-Layout — ohne
   2-Spalten-Tabelle rutscht der Empfaenger unter den Infoblock. Kostete rund
   7 Zeilen.
2. **Fusszeile im Fliesstext.** Der Footer war ein normaler Absatz am Textende
   und rutschte als einzige Zeile auf Seite 2. Gehoert in die echte
   Word-Seitenfusszeile (``section.footer``).

Referenz ist das Rechnungs-DOCX (``app/invoices/document_service.py``), das
genau dieses Layout schon hatte.
"""
import io
from datetime import date

import pytest
from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph

from app.extensions import db
from app.models import AppSetting, Customer, DunningNotice, Invoice

PT_PER_CM = 28.3465
# A4-Hoehe minus der im Generator gesetzten Raender (je 2 cm).
USABLE_CM = 29.7 - 2 - 2


@pytest.fixture
def notice(app):
    for key, val in [
        ("wg.name", "WG Treffling"),
        ("wg.address", "Treffling 7\n9871 Seeboden"),
        ("wg.email", "buero@treffling.test"),
        ("wg.phone", "12121212"),
        ("wg.iban", "AT12 3456 7890 1234 5678"),
        ("invoice.sender_address", "WG Treffling - Treffling 7 - 9871 Seeboden"),
    ]:
        AppSetting.set(key, val)
    c = Customer(name="Leeb Marco", first_name="Marco", last_name="Leeb",
                 strasse="Treffling", hausnummer="220", plz="9871",
                 ort="Seeboden am Millstätter See", customer_number="103")
    db.session.add(c)
    db.session.flush()
    inv = Invoice(customer_id=c.id, invoice_number="2026-00221",
                  date=date(2026, 6, 7), due_date=date(2026, 7, 7),
                  total_amount=312.61)
    db.session.add(inv)
    db.session.flush()
    n = DunningNotice(invoice_id=inv.id, level_snapshot=1,
                      name_snapshot="Zahlungserinnerung",
                      print_title_snapshot="Zahlungserinnerung",
                      issued_date=date(2026, 7, 21),
                      new_due_date=date(2026, 8, 4), fee_amount=0, status="Aktiv")
    db.session.add(n)
    db.session.commit()
    return n


def _build(notice):
    from app.dunning.document_service import generate_dunning_docx
    from app.settings_service import wg_settings
    return Document(io.BytesIO(generate_dunning_docx(notice, wg_settings())))


def _para_height_pt(p, default_pt=11.0):
    """Grobe Zeilenhoehe: groesste Schrift * 1.2 + Absatzabstaende."""
    sizes = [r.font.size.pt for r in p.runs if r.font.size]
    size = max(sizes) if sizes else default_pt
    pf = p.paragraph_format
    before = pf.space_before.pt if pf.space_before is not None else 0.0
    after = pf.space_after.pt if pf.space_after is not None else 10.0  # Word-Default
    spacing = pf.line_spacing if isinstance(pf.line_spacing, float) else 1.15
    return size * 1.2 * spacing + before + after


def _content_height_cm(doc):
    total = 0.0
    for child in doc.element.body.iterchildren():
        if child.tag.endswith("}p"):
            total += _para_height_pt(Paragraph(child, doc))
        elif child.tag.endswith("}tbl"):
            for row in Table(child, doc).rows:
                heights = [sum(_para_height_pt(p, 10.0) for p in cell.paragraphs)
                           for cell in row.cells]
                total += max(heights) if heights else 0.0
    return total / PT_PER_CM


class TestOnePage:
    def test_content_fits_on_one_page(self, notice):
        height = _content_height_cm(_build(notice))
        assert height < USABLE_CM, (
            f"Mahnung braucht ~{height:.1f} cm, auf die Seite passen "
            f"{USABLE_CM:.1f} cm -> Word bricht um."
        )

    def test_footer_is_a_real_page_footer(self, notice):
        doc = _build(notice)
        assert "WG Treffling" in doc.sections[0].footer.paragraphs[0].text
        # ... und NICHT zusaetzlich als letzter Absatz im Fliesstext.
        assert not [p for p in doc.paragraphs if "Treffling 7 |" in p.text]


class TestAddressBlock:
    def test_recipient_and_meta_are_side_by_side(self, notice):
        """Empfaenger links, Infoblock rechts — in EINER Tabellenzeile."""
        doc = _build(notice)
        matches = []
        for tbl in doc.tables:
            left = " ".join(p.text for p in tbl.cell(0, 0).paragraphs)
            right = " ".join(p.text for p in tbl.cell(0, 1).paragraphs) \
                if len(tbl.columns) > 1 else ""
            if "Marco Leeb" in left and "Rechnungsnummer" in right:
                matches.append(tbl)
        assert matches, "Empfaenger und Infoblock stehen nicht nebeneinander"

    def test_sender_return_address_sits_above_recipient(self, notice):
        """Der Absender gehoert nach oben, gegenueber dem Infoblock."""
        doc = _build(notice)
        for tbl in doc.tables:
            texts = [p.text for p in tbl.cell(0, 0).paragraphs]
            if any("Marco Leeb" in t for t in texts):
                joined = " | ".join(texts)
                assert "WG Treffling - Treffling 7" in joined, joined
                # Rueckadresse steht VOR dem Empfaengernamen.
                idx_sender = next(i for i, t in enumerate(texts) if "WG Treffling -" in t)
                idx_cust = next(i for i, t in enumerate(texts) if "Marco Leeb" in t)
                assert idx_sender < idx_cust
                return
        pytest.fail("Adressblock nicht gefunden")

    def test_meta_contains_all_dunning_fields(self, notice):
        doc = _build(notice)
        meta = ""
        for tbl in doc.tables:
            if len(tbl.columns) > 1:
                right = " ".join(p.text for p in tbl.cell(0, 1).paragraphs)
                if "Rechnungsnummer" in right:
                    meta = right
        for label in ("Rechnungsnummer", "Kundennummer", "Rechnungsdatum",
                      "Mahndatum", "Zahlbar bis"):
            assert label in meta, f"{label} fehlt im Infoblock"
