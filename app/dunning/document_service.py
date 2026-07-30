"""Hilfsfunktionen zur Mahn-Dokument-Generierung (ADR-003)."""
import io
from decimal import Decimal

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, Cm, RGBColor

from app.invoices.design import get_design


def _hex_to_rgb(value: str) -> RGBColor:
    v = (value or "").lstrip("#")
    if len(v) != 6:
        return RGBColor(0x33, 0x33, 0x33)
    return RGBColor(int(v[0:2], 16), int(v[2:4], 16), int(v[4:6], 16))


def _hex_fill(value: str) -> str:
    return (value or "").lstrip("#").upper() or "FFFFFF"


def _right_align_cell(cell, text: str, *, font_name: str | None = None,
                      font_size=None, color: RGBColor | None = None,
                      bold: bool = False) -> None:
    """Setzt Text + RIGHT-Alignment sicher (Alignment NACH Text)."""
    cell.text = text or ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    if p.runs:
        run = p.runs[0]
        if font_name:
            run.font.name = font_name
        if font_size is not None:
            run.font.size = font_size
        if color is not None:
            run.font.color.rgb = color
        if bold:
            run.bold = True


# ---------------------------------------------------------------------------
# Helpers (identisch zu app/invoices/document_service.py)
# ---------------------------------------------------------------------------

def _de_fmt(value, decimals=2) -> str:
    try:
        value = Decimal(str(value))
    except Exception:
        return str(value)
    formatted = f"{value:,.{decimals}f}"
    return formatted.replace(",", "X").replace(".", ",").replace("X", ".")


def _set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _set_cell_border_bottom(cell, size_pt: int = 12, color_hex: str = "333333"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size_pt))
    bottom.set(qn("w:space"), "0")
    bottom.set(qn("w:color"), _hex_fill(color_hex))
    tcBorders.append(bottom)
    tcPr.append(tcBorders)


def _tight(paragraph, *, before: int = 0, after: int = 2):
    """Absatz eng setzen (Adress-/Meta-/Infoblöcke).

    Word gibt jedem Absatz standardmäßig ~10 pt Abstand nach unten. In den
    mehrzeiligen Blöcken summiert sich das auf mehrere Zentimeter und war der
    zweite Grund für den Seitenumbruch.
    """
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = 1.0
    return paragraph


def _remove_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "none")
        tblBorders.append(el)
    tblPr.append(tblBorders)


# ---------------------------------------------------------------------------
# DOCX-Generierung
# ---------------------------------------------------------------------------

def generate_dunning_docx(notice, wg: dict, design: dict | None = None,
                          invoice_sender_address: str | None = None) -> bytes:
    """Erstellt ein Word-Dokument (.docx) für eine Mahnung.

    Layout wie beim Rechnungs-DOCX (``app/invoices/document_service.py``):
    Briefkopf, darunter **Absender-Rückadresse + Empfänger links, Infoblock
    rechts** als rahmenlose 2-Spalten-Tabelle. Word kennt kein ``float`` wie das
    PDF-Layout — ohne Tabelle rutscht der Empfänger unter den Infoblock statt
    daneben. Genau das war der Fehler: die gestapelten Blöcke kosteten rund
    7 Zeilen Höhe und schoben die Fußzeile auf eine zweite Seite.

    Parameters
    ----------
    notice : DunningNotice
        Die Mahnung mit verknüpfter Invoice.
    wg : dict
        WG-Kontaktdaten (name, address, email, phone, iban, bic).
    design : dict | None
        Design-Parameter (Schriftart, Farben). Wenn ``None``, wird das
        Standard-Design ``classic`` verwendet.
    invoice_sender_address : str | None
        Einzeilige Absender-Rückadresse über der Empfängeranschrift
        (Fensterkuvert). ``None`` lädt den gespeicherten Wert.

    Returns
    -------
    bytes
        Rohe .docx-Datei als Bytes.
    """
    from app.dunning.services import dunning_summary, rendered_letter_texts

    if design is None:
        design = get_design("classic")
    if invoice_sender_address is None:
        from app.settings_service import get_invoice_sender_address
        invoice_sender_address = get_invoice_sender_address()

    font_name = design.get("docx_font", "Arial")
    text_rgb = _hex_to_rgb(design.get("text_color", "#333333"))
    muted_rgb = _hex_to_rgb(design.get("muted_color", "#666666"))
    heading_rgb = _hex_to_rgb(design.get("heading_color", "#333333"))
    accent_rgb = _hex_to_rgb(design.get("accent_color", "#333333"))
    rule_hex = _hex_fill(design.get("rule_color", "#333333"))
    header_bg_hex = _hex_fill(design.get("header_bg", "#F0F0F0"))
    header_text_rgb = _hex_to_rgb(design.get("header_text", "#333333"))
    payment_bg_hex = _hex_fill(design.get("payment_bg", "#F9F9F9"))

    invoice = notice.invoice
    customer = invoice.customer
    summary = dunning_summary(invoice)
    intro_text, closing_text = rendered_letter_texts(notice, summary, wg)

    doc = Document()

    # ── Seitenränder ─────────────────────────────────────────────────────
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # ── Standardschrift ───────────────────────────────────────────────────
    style = doc.styles["Normal"]
    style.font.name = font_name
    style.font.size = Pt(11)
    style.font.color.rgb = text_rgb

    # ── Briefkopf (2-Spalten-Tabelle ohne Rahmen) ────────────────────────
    header_tbl = doc.add_table(rows=1, cols=2)
    _remove_table_borders(header_tbl)
    header_tbl.columns[0].width = Cm(10)
    header_tbl.columns[1].width = Cm(6)

    left_cell = header_tbl.cell(0, 0)
    p_name = left_cell.paragraphs[0]
    p_name.paragraph_format.space_after = Pt(0)
    run_name = p_name.add_run(wg.get("name", ""))
    run_name.bold = True
    run_name.font.size = Pt(12)
    run_name.font.color.rgb = heading_rgb

    address = wg.get("address", "")
    if address:
        for line in address.replace("\\n", "\n").split("\n"):
            p_addr = left_cell.add_paragraph(line.strip())
            _tight(p_addr)
            p_addr.runs[0].font.size = Pt(9)
            p_addr.runs[0].font.color.rgb = muted_rgb

    right_cell = header_tbl.cell(0, 1)
    right_cell.paragraphs[0].clear()
    for key in ("email", "phone"):
        val = wg.get(key, "")
        if val:
            p = right_cell.add_paragraph(val)
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            _tight(p)
            p.runs[0].font.size = Pt(9)
            p.runs[0].font.color.rgb = accent_rgb
    if not right_cell.paragraphs[0].text:
        right_cell.paragraphs[0]._element.getparent().remove(
            right_cell.paragraphs[0]._element
        )

    doc.add_paragraph()  # Abstand

    # ── Absender + Empfänger (links) | Infoblock (rechts) ────────────────
    meta_lines = [
        ("Rechnungsnummer", invoice.invoice_number),
    ]
    if customer.customer_number:
        meta_lines.append(("Kundennummer", str(customer.customer_number)))
    meta_lines.append(("Rechnungsdatum", invoice.date.strftime("%d.%m.%Y")))
    meta_lines.append(("Mahndatum", notice.issued_date.strftime("%d.%m.%Y")))
    if notice.new_due_date:
        meta_lines.append(("Zahlbar bis", notice.new_due_date.strftime("%d.%m.%Y")))

    addr_tbl = doc.add_table(rows=1, cols=2)
    _remove_table_borders(addr_tbl)
    addr_tbl.columns[0].width = Cm(9)
    addr_tbl.columns[1].width = Cm(7)
    recipient_cell = addr_tbl.cell(0, 0)
    meta_cell = addr_tbl.cell(0, 1)

    # Absender-Rückadresse (klein, Fensterkuvert) über der Anschrift.
    p_cust = recipient_cell.paragraphs[0]
    if invoice_sender_address:
        run_ret = p_cust.add_run(invoice_sender_address)
        run_ret.font.size = Pt(7)
        run_ret.font.color.rgb = muted_rgb
        p_cust.paragraph_format.space_after = Pt(0)
        p_cust = recipient_cell.add_paragraph()
    p_cust.add_run(customer.letter_name).bold = True
    _tight(p_cust, after=0)

    street_parts = [customer.strasse, customer.hausnummer]
    street = " ".join(p for p in street_parts if p)
    city_parts = [customer.plz, customer.ort]
    city = " ".join(p for p in city_parts if p)
    if street:
        _tight(recipient_cell.add_paragraph(street), after=0)
    if city:
        _tight(recipient_cell.add_paragraph(city), after=0)
    land = customer.land
    if land and land != "Österreich":
        _tight(recipient_cell.add_paragraph(land), after=0)

    # Infoblock rechts, auf gleicher Höhe wie der Absender.
    meta_cell.paragraphs[0].clear()
    meta_first = True
    for label, value in meta_lines:
        p = meta_cell.paragraphs[0] if meta_first else meta_cell.add_paragraph()
        meta_first = False
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        _tight(p)
        run_lbl = p.add_run(f"{label}: ")
        run_lbl.bold = True
        run_lbl.font.color.rgb = heading_rgb
        p.add_run(value)

    # ── Überschrift ──────────────────────────────────────────────────────
    title = notice.print_title_snapshot or notice.name_snapshot
    heading = doc.add_heading(title, level=1)
    heading.paragraph_format.space_before = Pt(18)
    heading.paragraph_format.space_after = Pt(6)
    heading.runs[0].font.color.rgb = heading_rgb
    heading.runs[0].font.name = font_name

    # ── Einleitungstext (pro Stufe konfigurierbar, siehe services) ───────
    doc.add_paragraph("Sehr geehrte Damen und Herren,")
    for para in (intro_text or "").split("\n"):
        if para.strip():
            doc.add_paragraph(para.strip())

    # ── Forderungsübersicht ──────────────────────────────────────────────
    tbl = doc.add_table(rows=1, cols=2)
    tbl.style = "Table Grid"
    tbl.columns[0].width = Cm(11)
    tbl.columns[1].width = Cm(5)

    hdr = tbl.rows[0].cells
    hdr[0].text = "Position"
    _right_align_cell(hdr[1], "Betrag", font_name=font_name,
                      font_size=Pt(10), color=header_text_rgb, bold=True)
    # Position-Kopf ohne Ausrichtung links — nur Formatierung
    run_pos = hdr[0].paragraphs[0].runs[0]
    run_pos.bold = True
    run_pos.font.size = Pt(10)
    run_pos.font.name = font_name
    run_pos.font.color.rgb = header_text_rgb
    for i in range(2):
        _set_cell_bg(hdr[i], header_bg_hex)
        _set_cell_border_bottom(hdr[i], 16, rule_hex)

    # Hauptforderung (ursprünglicher Rechnungsbetrag)
    row = tbl.add_row().cells
    row[0].text = f"Rechnungsbetrag ({invoice.invoice_number})"
    _right_align_cell(row[1], f"{_de_fmt(summary['original_total'], 2)} €", font_size=Pt(10))
    for r in row[0].paragraphs[0].runs:
        r.font.size = Pt(10)

    # Bereits geleistete Zahlung (nur bei Teilzahlung)
    if summary.get("paid") and summary["paid"] > 0:
        row = tbl.add_row().cells
        row[0].text = "abzüglich Zahlungseingang"
        _right_align_cell(row[1], f"− {_de_fmt(summary['paid'], 2)} €", font_size=Pt(10))
        for r in row[0].paragraphs[0].runs:
            r.font.size = Pt(10)

    # Kumulative Mahngebühren
    for n in summary["notices"]:
        if n.fee_amount and n.fee_amount > 0:
            row = tbl.add_row().cells
            row[0].text = f"Mahngebühr – {n.name_snapshot}"
            _right_align_cell(row[1], f"{_de_fmt(n.fee_amount, 2)} €", font_size=Pt(10))
            for r in row[0].paragraphs[0].runs:
                r.font.size = Pt(10)

    # Gesamtbetrag
    row_total = tbl.add_row().cells
    row_total[0].text = "Gesamtbetrag"
    run_total_lbl = row_total[0].paragraphs[0].runs[0]
    run_total_lbl.bold = True
    run_total_lbl.font.color.rgb = heading_rgb
    _right_align_cell(row_total[1], f"{_de_fmt(summary['gross_total'], 2)} €",
                      color=heading_rgb, bold=True)

    doc.add_paragraph()

    # ── Zahlungsaufforderung ─────────────────────────────────────────────
    new_due_str = notice.new_due_date.strftime("%d.%m.%Y") if notice.new_due_date else "—"

    payment_tbl = doc.add_table(rows=1, cols=1)
    payment_tbl.style = "Table Grid"
    payment_cell = payment_tbl.cell(0, 0)
    _set_cell_bg(payment_cell, payment_bg_hex)

    p_pay = _tight(payment_cell.paragraphs[0])
    run_pay_lbl = p_pay.add_run("Zahlungsinformationen")
    run_pay_lbl.bold = True
    run_pay_lbl.font.color.rgb = heading_rgb

    p_pay2 = _tight(payment_cell.add_paragraph("Bitte überweisen Sie den Betrag von "))
    p_pay2.add_run(f"{_de_fmt(summary['gross_total'], 2)} €").bold = True
    p_pay2.add_run(f" bis zum {new_due_str}")

    if wg.get("iban"):
        p_iban = _tight(payment_cell.add_paragraph("IBAN: "))
        p_iban.add_run(wg["iban"]).bold = True
    if wg.get("bic"):
        _tight(payment_cell.add_paragraph(f"BIC: {wg['bic']}"))
    _tight(payment_cell.add_paragraph(
        f"Empfänger: {wg.get('account_holder') or wg.get('name', '')}"
    ))
    p_ref = _tight(payment_cell.add_paragraph("Verwendungszweck: "))
    p_ref.add_run(f"{invoice.invoice_number} / Mahnung").bold = True

    doc.add_paragraph()

    # ── Schlusstext (pro Stufe konfigurierbar, siehe services) ───────────
    for para in (closing_text or "").split("\n"):
        if para.strip():
            doc.add_paragraph(para.strip())

    # ── Fußzeile: echte Word-Seitenfußzeile ──────────────────────────────
    # Frueher ein normaler Absatz am Textende — der rutschte bei gut gefuelltem
    # Blatt als einzige Zeile auf eine zweite Seite (genau der gemeldete Fehler).
    # In der Seitenfusszeile liegt er im unteren Seitenrand und kann per
    # Definition keinen Umbruch ausloesen; zusaetzlich steht er dann auf JEDER
    # Seite, falls eine Mahnung doch einmal laenger wird.
    footer_parts = [wg.get("name", "")]
    addr = wg.get("address", "")
    if addr:
        footer_parts.append(addr.replace("\\n", " | ").replace("\n", " | "))
    if wg.get("email"):
        footer_parts.append(wg["email"])
    p_footer = section.footer.paragraphs[0]
    p_footer.text = " \u2014 ".join(p for p in footer_parts if p)
    p_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _tight(p_footer)
    if p_footer.runs:
        p_footer.runs[0].font.name = font_name
        p_footer.runs[0].font.size = Pt(8)
        p_footer.runs[0].font.color.rgb = muted_rgb

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
