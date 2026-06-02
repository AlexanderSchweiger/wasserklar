"""Hilfsfunktionen zur Rechnungsdokument-Generierung."""
import io
import re
from decimal import Decimal
from html.parser import HTMLParser

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, Cm, RGBColor

from app.invoices.design import get_design


def _hex_to_rgb(value: str) -> RGBColor:
    v = (value or "").lstrip("#")
    if len(v) != 6:
        return RGBColor(0x33, 0x33, 0x33)
    return RGBColor(int(v[0:2], 16), int(v[2:4], 16), int(v[4:6], 16))


def _hex_fill(value: str) -> str:
    return (value or "").lstrip("#").upper() or "FFFFFF"


def _right_align_cell(cell, text: str, *, font_name: str | None = None,
                      font_size=None, color: RGBColor | None = None,
                      bold: bool = False) -> None:
    """Setzt Text + RIGHT-Alignment sicher (Alignment NACH Text, damit es
    auch bei älteren python-docx-Versionen robust bleibt).

    Optional werden Schrift-Eigenschaften am ersten Run gesetzt.
    """
    cell.text = text or ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    if p.runs:
        run = p.runs[0]
        if font_name:
            run.font.name = font_name
        if font_size is not None:
            run.font.size = font_size
        if color is not None:
            run.font.color.rgb = color
        if bold:
            run.bold = True


def _de_fmt(value, decimals=2) -> str:
    """Formatiert eine Zahl im deutschen Format (1.250,90)."""
    try:
        value = Decimal(str(value))
    except Exception:
        return str(value)
    formatted = f"{value:,.{decimals}f}"
    # Python nutzt englisches Format: tausend=Komma, dezimal=Punkt → tauschen
    return formatted.replace(",", "X").replace(".", ",").replace("X", ".")


def _set_cell_bg(cell, hex_color: str):
    """Setzt die Hintergrundfarbe einer Tabellenzelle."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _set_cell_border_bottom(cell, size_pt: int = 12, color_hex: str = "333333"):
    """Setzt einen unteren Rahmen an einer Tabellenzelle."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size_pt))
    bottom.set(qn("w:space"), "0")
    bottom.set(qn("w:color"), _hex_fill(color_hex))
    tcBorders.append(bottom)
    tcPr.append(tcBorders)


def _remove_table_borders(table):
    """Entfernt alle Rahmenlinien einer Tabelle (für Briefkopf-Layout)."""
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "none")
        tblBorders.append(el)
    tblPr.append(tblBorders)


_FONT_SIZE_RE = re.compile(r'font-size\s*:\s*(\d+(?:\.\d+)?)\s*pt', re.IGNORECASE)


class _RichTextRunParser(HTMLParser):
    """Parst das Kontakttext-HTML in Zeilen von Runs:
    Liste von Zeilen, jede Zeile = Liste (text, bold, italic, underline, font_size_pt|None)."""

    def __init__(self):
        super().__init__(convert_charrefs=True)
        self.lines = [[]]
        self._b = self._i = self._u = False
        self._font_sizes = []  # Stack für verschachtelte font-size-Spans

    @property
    def _font_size(self):
        return self._font_sizes[-1] if self._font_sizes else None

    def handle_starttag(self, tag, attrs):
        if tag == "b":
            self._b = True
        elif tag == "i":
            self._i = True
        elif tag == "u":
            self._u = True
        elif tag == "br":
            self.lines.append([])
        elif tag == "span":
            style = dict(attrs).get('style', '')
            m = _FONT_SIZE_RE.search(style)
            self._font_sizes.append(int(float(m.group(1))) if m else None)

    def handle_startendtag(self, tag, attrs):
        if tag == "br":
            self.lines.append([])

    def handle_endtag(self, tag):
        if tag == "b":
            self._b = False
        elif tag == "i":
            self._i = False
        elif tag == "u":
            self._u = False
        elif tag == "span" and self._font_sizes:
            self._font_sizes.pop()

    def handle_data(self, data):
        if data:
            self.lines[-1].append((data, self._b, self._i, self._u, self._font_size))


def _parse_rich_text(html: str) -> list:
    """Parst Kontakttext-HTML in Zeilen von formatierten Runs."""
    if not html:
        return []
    parser = _RichTextRunParser()
    parser.feed(html)
    return parser.lines


def generate_docx(invoice, wg: dict, design: dict | None = None,
                   contact_info: str | None = None,
                   contact_info_font_size: int | None = None,
                   invoice_sender_address: str | None = None) -> bytes:
    """Erstellt ein Word-Dokument (.docx) für die übergebene Rechnung.

    Parameters
    ----------
    invoice : Invoice
        Das Invoice-Objekt mit allen verknüpften Daten.
    wg : dict
        WG-Kontaktdaten (name, address, email, phone, iban, bic).
    design : dict | None
        Design-Parameter (Schriftart, Farben). Wenn ``None``, wird das
        Standard-Design ``classic`` verwendet.
    contact_info : str | None
        Sanitisiertes Kontakttext-HTML (<b>/<i>/<u>/<br>) für den Adressblock.
        Wenn ``None``, wird der gespeicherte Wert aus den AppSettings geladen.

    Returns
    -------
    bytes
        Rohe .docx-Datei als Bytes.
    """
    if design is None:
        design = get_design("classic")
    if contact_info is None:
        from app.settings_service import get_contact_info
        contact_info = get_contact_info()
    if contact_info_font_size is None:
        from app.settings_service import get_contact_info_font_size
        contact_info_font_size = get_contact_info_font_size()
    if invoice_sender_address is None:
        from app.settings_service import get_invoice_sender_address
        invoice_sender_address = get_invoice_sender_address()

    font_name = design.get("docx_font", "Arial")
    text_rgb = _hex_to_rgb(design.get("text_color", "#333333"))
    muted_rgb = _hex_to_rgb(design.get("muted_color", "#666666"))
    heading_rgb = _hex_to_rgb(design.get("heading_color", "#333333"))
    rule_hex = _hex_fill(design.get("rule_color", "#333333"))
    header_bg_hex = _hex_fill(design.get("header_bg", "#F0F0F0"))
    payment_bg_hex = _hex_fill(design.get("payment_bg", "#F9F9F9"))

    doc = Document()

    # ── Seitenränder ─────────────────────────────────────────────────────
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # ── Standardschrift ───────────────────────────────────────────────────
    style = doc.styles["Normal"]
    style.font.name = font_name
    style.font.size = Pt(11)
    style.font.color.rgb = text_rgb

    # ── Briefkopf (2-Spalten-Tabelle ohne Rahmen) ─────────────────────────
    header_tbl = doc.add_table(rows=1, cols=2)
    _remove_table_borders(header_tbl)
    header_tbl.columns[0].width = Cm(10)
    header_tbl.columns[1].width = Cm(6)

    # Linke Spalte bleibt leer: die WG-Identitaet (Name/Adresse/E-Mail/Telefon)
    # wird nicht mehr automatisch gesetzt — der Briefkopf besteht jetzt allein aus
    # dem frei gestaltbaren Kontaktinfo-Block (rechts).
    left_cell = header_tbl.cell(0, 0)
    left_cell.paragraphs[0].clear()

    right_cell = header_tbl.cell(0, 1)
    right_cell.paragraphs[0].clear()
    rt_first = True
    for line in _parse_rich_text(contact_info):
        p = right_cell.paragraphs[0] if rt_first else right_cell.add_paragraph()
        rt_first = False
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        # Enger Zeilenabstand im Kontaktblock (kein Absatzabstand, einfache Höhe).
        pf = p.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        pf.line_spacing = 1.0
        for text, bold, italic, underline, run_font_size in line:
            run = p.add_run(text)
            run.font.size = Pt(run_font_size if run_font_size is not None else contact_info_font_size)
            run.bold = bold
            run.italic = italic
            run.underline = underline
    # Leere erste Zeile in right_cell entfernen
    if not right_cell.paragraphs[0].text:
        right_cell.paragraphs[0]._element.getparent().remove(right_cell.paragraphs[0]._element)

    doc.add_paragraph()  # Abstand

    # ── Empfänger (links) + Rechnungs-Meta (rechts) ───────────────────────
    # Als rahmenlose 2-Spalten-Tabelle: Word kennt kein float wie das
    # PDF-Layout — ohne Tabelle rutscht der Empfängerblock unter den
    # gesamten Meta-Block statt daneben zu stehen.
    invoice_date_str = invoice.date.strftime("%d.%m.%Y")
    meta_lines = [
        ("Rechnungsnummer", invoice.invoice_number),
        ("Rechnungsdatum", invoice_date_str),
        ("Lieferdatum", invoice_date_str),
    ]
    if invoice.customer.customer_number:
        meta_lines.append(("Kundennummer", str(invoice.customer.customer_number)))
    meta_lines.append((
        "Fällig bis",
        invoice.due_date.strftime("%d.%m.%Y") if invoice.due_date else "—"
    ))
    if invoice.billing_period:
        meta_lines.append(("Abrechnungsperiode", invoice.billing_period.name))

    addr_tbl = doc.add_table(rows=1, cols=2)
    _remove_table_borders(addr_tbl)
    addr_tbl.columns[0].width = Cm(9)
    addr_tbl.columns[1].width = Cm(7)
    recipient_cell = addr_tbl.cell(0, 0)
    meta_cell = addr_tbl.cell(0, 1)

    # Empfänger (linke Spalte)
    p_cust = recipient_cell.paragraphs[0]
    if invoice_sender_address:
        run_ret = p_cust.add_run(invoice_sender_address)
        run_ret.font.size = Pt(7)
        run_ret.font.color.rgb = muted_rgb
        p_cust.paragraph_format.space_after = Pt(0)
        p_cust = recipient_cell.add_paragraph()
    p_cust.add_run(invoice.customer.name).bold = True

    street_parts = [invoice.customer.strasse, invoice.customer.hausnummer]
    street = " ".join(p for p in street_parts if p)
    city_parts = [invoice.customer.plz, invoice.customer.ort]
    city = " ".join(p for p in city_parts if p)
    if street:
        recipient_cell.add_paragraph(street)
    if city:
        recipient_cell.add_paragraph(city)
    land = invoice.customer.land
    if land and land != "Österreich":
        recipient_cell.add_paragraph(land)
    if invoice.property:
        p_prop = recipient_cell.add_paragraph()
        run_prop = p_prop.add_run(f"Betr.: {invoice.property.label()}")
        run_prop.font.size = Pt(9)

    # Rechnungs-Meta + Kontaktinfo (rechte Spalte, rechtsbündig)
    meta_cell.paragraphs[0].clear()
    meta_first = True
    for label, value in meta_lines:
        p = meta_cell.paragraphs[0] if meta_first else meta_cell.add_paragraph()
        meta_first = False
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run_lbl = p.add_run(f"{label}: ")
        run_lbl.bold = True
        run_lbl.font.color.rgb = heading_rgb
        p.add_run(value)

    # ── Überschrift ───────────────────────────────────────────────────────
    heading = doc.add_heading(f"Rechnung {invoice.invoice_number}", level=1)
    heading.runs[0].font.color.rgb = heading_rgb
    heading.runs[0].font.name = font_name

    # ── Einleitungstext ───────────────────────────────────────────────────
    if invoice.property:
        p_intro = doc.add_paragraph()
        p_intro.add_run("Wir stellen für das Objekt ")
        prop_display = invoice.property.address_display() or invoice.property.label()
        p_intro.add_run(f"„{prop_display}”").bold = True
        p_intro.add_run(" wie folgt in Rechnung:")
    else:
        doc.add_paragraph("Wir stellen wie folgt in Rechnung:")

    # ── Positionen-Tabelle ────────────────────────────────────────────────
    items = invoice.items
    tax_summary = invoice.tax_breakdown  # {rate: {"net", "tax"}}
    has_tax = bool(tax_summary)
    col_count = 5  # Beschreibung / Menge / Einheit / Einzelpreis / Betrag

    tbl = doc.add_table(rows=1, cols=col_count)
    tbl.style = "Table Grid"

    # Kopfzeile
    hdr_cells = tbl.rows[0].cells
    headers = ["Beschreibung", "Menge", "Einheit", "Einzelpreis", "Gesamtpreis"]

    header_text_rgb = _hex_to_rgb(design.get("header_text", "#333333"))
    for i, hdr in enumerate(headers):
        right = hdr in ("Menge", "Einheit", "Einzelpreis", "Gesamtpreis")
        if right:
            _right_align_cell(hdr_cells[i], hdr, font_name=font_name,
                              font_size=Pt(10), color=header_text_rgb, bold=True)
        else:
            hdr_cells[i].text = hdr
            run = hdr_cells[i].paragraphs[0].runs[0]
            run.bold = True
            run.font.size = Pt(10)
            run.font.name = font_name
            run.font.color.rgb = header_text_rgb
        _set_cell_bg(hdr_cells[i], header_bg_hex)
        _set_cell_border_bottom(hdr_cells[i], 16, rule_hex)

    # Positionen
    for item in items:
        row_cells = tbl.add_row().cells
        row_cells[0].text = item.description or ""
        _right_align_cell(row_cells[1], _de_fmt(item.quantity, 2), font_size=Pt(10))
        _right_align_cell(row_cells[2], item.unit or "", font_size=Pt(10))
        _right_align_cell(row_cells[3], f"{_de_fmt(item.unit_price, 2)} €", font_size=Pt(10))
        _right_align_cell(row_cells[4], f"{_de_fmt(item.amount, 2)} €", font_size=Pt(10))
        for run in row_cells[0].paragraphs[0].runs:
            run.font.size = Pt(10)

    # Summenzeilen
    if has_tax:
        # Nettosumme
        row_net = tbl.add_row().cells
        row_net[0].merge(row_net[col_count - 2])
        _right_align_cell(row_net[0], "Nettobetrag")
        _right_align_cell(row_net[col_count - 1], f"{_de_fmt(invoice.net_total, 2)} €")
        # USt pro Satz
        for rate, info in tax_summary.items():
            row_tax = tbl.add_row().cells
            row_tax[0].merge(row_tax[col_count - 2])
            _right_align_cell(
                row_tax[0],
                f"zzgl. USt. {_de_fmt(rate, 0)} % auf {_de_fmt(info['net'], 2)} €",
            )
            _right_align_cell(row_tax[col_count - 1], f"{_de_fmt(info['tax'], 2)} €")
        # Gesamt
        row_total = tbl.add_row().cells
        row_total[0].merge(row_total[col_count - 2])
        _right_align_cell(row_total[0], "Gesamtsumme",
                          color=heading_rgb, bold=True)
        _right_align_cell(row_total[col_count - 1],
                          f"{_de_fmt(invoice.total_amount, 2)} €",
                          color=heading_rgb, bold=True)
    else:
        _add_total_row(tbl, col_count, invoice.total_amount, "Gesamtsumme", heading_rgb)

    # ── Hinweistext ───────────────────────────────────────────────────────
    if invoice.notes:
        p_notes = doc.add_paragraph()
        run_notes = p_notes.add_run(invoice.notes)
        run_notes.italic = True

    doc.add_paragraph()  # Abstand

    # ── Zahlungsinformationen ─────────────────────────────────────────────
    payment_tbl = doc.add_table(rows=1, cols=1)
    payment_tbl.style = "Table Grid"
    payment_cell = payment_tbl.cell(0, 0)
    _set_cell_bg(payment_cell, payment_bg_hex)

    p_pay = payment_cell.paragraphs[0]
    run_pay_lbl = p_pay.add_run("Zahlung")
    run_pay_lbl.bold = True
    run_pay_lbl.font.color.rgb = heading_rgb

    p_pay2 = payment_cell.add_paragraph("Wir ersuchen Sie, den Rechnungsbetrag von ")
    p_pay2.add_run(f"{_de_fmt(invoice.total_amount, 2)} €").bold = True
    if invoice.due_date:
        p_pay2.add_run(f" bis zum {invoice.due_date.strftime('%d.%m.%Y')}")
    p_pay2.add_run(" auf unser Konto einzuzahlen:")

    if wg.get("iban"):
        p_iban = payment_cell.add_paragraph("IBAN: ")
        p_iban.add_run(wg["iban"]).bold = True
    if wg.get("bic"):
        p_bic = payment_cell.add_paragraph("BIC: ")
        p_bic.add_run(wg["bic"]).bold = True
    payment_cell.add_paragraph(
        f"Empfänger: {wg.get('account_holder') or wg.get('name', '')}")
    p_ref = payment_cell.add_paragraph("Verwendungszweck: ")
    p_ref.add_run(invoice.invoice_number).bold = True

    doc.add_paragraph()  # Abstand

    # ── Grußformel ────────────────────────────────────────────────────────
    doc.add_paragraph("Mit freundlichen Grüßen")
    doc.add_paragraph(wg.get("name", ""))

    doc.add_paragraph()  # Abstand

    # ── Fußzeile ──────────────────────────────────────────────────────────
    footer_parts = [wg.get("name", "")]
    addr = wg.get("address", "")
    if addr:
        footer_parts.append(addr.replace("\\n", " | ").replace("\n", " | "))
    if wg.get("email"):
        footer_parts.append(wg["email"])
    p_footer = doc.add_paragraph(" \u2014 ".join(footer_parts))
    p_footer.runs[0].font.size = Pt(9)
    p_footer.runs[0].font.color.rgb = muted_rgb

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _add_total_row(tbl, col_count: int, amount, label: str, color_rgb: RGBColor | None = None):
    row = tbl.add_row().cells
    row[0].merge(row[col_count - 2])
    _right_align_cell(row[0], label, color=color_rgb, bold=True)
    _right_align_cell(row[col_count - 1], f"{_de_fmt(amount, 2)} €",
                      color=color_rgb, bold=True)


def merge_docx_files(sources: list) -> bytes:
    """Fügt mehrere .docx-Dateien zu einem Dokument zusammen.

    Parameters
    ----------
    sources : list
        Liste von Dateipfaden (str) oder Byte-Inhalten (bytes) der Quelldokumente.

    Returns
    -------
    bytes
        Das zusammengeführte .docx als Bytes.
    """
    from copy import deepcopy

    merged = Document()
    # Leeres Standard-Dokument bereinigen
    for el in list(merged.element.body):
        merged.element.body.remove(el)

    for idx, source in enumerate(sources):
        if idx > 0:
            # Seitenumbruch zwischen Rechnungen
            p = OxmlElement("w:p")
            r = OxmlElement("w:r")
            br = OxmlElement("w:br")
            br.set(qn("w:type"), "page")
            r.append(br)
            p.append(r)
            merged.element.body.append(p)

        if isinstance(source, (str, bytes)):
            src_doc = Document(io.BytesIO(source) if isinstance(source, bytes) else source)
        else:
            src_doc = Document(source)

        for element in src_doc.element.body:
            merged.element.body.append(deepcopy(element))

    buf = io.BytesIO()
    merged.save(buf)
    return buf.getvalue()
