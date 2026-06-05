"""PDF-/Word-Erzeugung für Einladungen und Protokolle.

PDF läuft über WeasyPrint (wie bei den Rechnungen — Import wird vom Aufrufer
mit ImportError/OSError-Fallback gekapselt); Word über python-docx. Das
Rechnungs-Design (Farben/Schrift, Briefkopf) wird wiederverwendet.
"""
import io
import re
from html import unescape

from flask import render_template

from app.models import AppSetting
from app.invoices.design import get_design
from app.settings_service import (
    wg_settings, get_contact_info, get_contact_info_font_size,
    get_invoice_sender_address,
)


def current_design():
    return get_design(AppSetting.get("invoice.design", "classic"))


# ── PDF (WeasyPrint-HTML) ────────────────────────────────────────────────────

def render_invitation_html(meeting, customer, agenda_items):
    """HTML einer einzelnen Einladung (für WeasyPrint), im Rechnungs-Stil."""
    return render_template(
        "schriftfuehrung/pdf/invitation.html",
        meeting=meeting, customer=customer, agenda_items=agenda_items,
        design=current_design(),
        contact_info=get_contact_info(),
        contact_info_font_size=get_contact_info_font_size(),
        sender_address=get_invoice_sender_address(),
        wg=wg_settings(),
    )


def render_protocol_html(meeting, protocol, attendances, resolutions, quorum):
    """HTML des Protokolls (für WeasyPrint): Kopf mit Anwesenheit/Quorum,
    Beschluss-Tabelle und narrativer Rich-Text."""
    return render_template(
        "schriftfuehrung/pdf/protocol.html",
        meeting=meeting, protocol=protocol, attendances=attendances,
        resolutions=resolutions, quorum=quorum,
        design=current_design(),
        contact_info=get_contact_info(),
        contact_info_font_size=get_contact_info_font_size(),
        wg=wg_settings(),
    )


# ── Word (python-docx) ───────────────────────────────────────────────────────

_BR_RE = re.compile(r"<br\s*/?>", re.IGNORECASE)
_BLOCK_RE = re.compile(r"</(p|div|li)>", re.IGNORECASE)
_TAG_RE = re.compile(r"<[^>]+>")


def html_to_text(html):
    """Wandelt das reduzierte Rich-Text-HTML (b/i/u/br/span) in reinen Text mit
    Zeilenumbrüchen — für die Word-Ausgabe (formatfrei, aber lesbar)."""
    if not html:
        return ""
    text = _BR_RE.sub("\n", html)
    text = _BLOCK_RE.sub("\n", text)
    text = _TAG_RE.sub("", text)
    return unescape(text).strip()


def _meeting_when(meeting):
    parts = []
    if meeting.meeting_date:
        parts.append(meeting.meeting_date.strftime("%d.%m.%Y"))
    if meeting.start_time:
        t = meeting.start_time.strftime("%H:%M")
        if meeting.end_time:
            t += "–" + meeting.end_time.strftime("%H:%M")
        parts.append(t + " Uhr")
    return ", ".join(parts)


def build_invitation_docx(meeting, customer, agenda_items, type_label):
    """Erzeugt eine Einladung als .docx (bytes) — ein schlichter Brief mit
    Briefkopf, Empfänger, Agenda und Schlusstext."""
    from docx import Document
    from docx.shared import Pt, Cm

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2)

    wg = wg_settings()

    # Briefkopf (Kontaktinfo als Text, rechtsbündig).
    header = html_to_text(get_contact_info()) or (wg.get("name") or "")
    if header:
        p = doc.add_paragraph()
        p.alignment = 2  # rechts
        run = p.add_run(header)
        run.font.size = Pt(9)

    # Empfänger-Adressblock.
    doc.add_paragraph()
    addr = doc.add_paragraph()
    addr.add_run(customer.name).bold = True
    if customer.address_display():
        addr.add_run("\n" + customer.address_display())

    # Betreff.
    doc.add_paragraph()
    subj = doc.add_paragraph()
    subj.add_run(f"Einladung zur {type_label}").bold = True
    when = _meeting_when(meeting)
    meta = []
    if when:
        meta.append(when)
    if meeting.location:
        meta.append(meeting.location)
    if meta:
        doc.add_paragraph(" · ".join(meta))

    # Einleitung.
    intro = html_to_text(meeting.intro_text)
    if intro:
        doc.add_paragraph(intro)

    # Tagesordnung.
    doc.add_paragraph()
    doc.add_paragraph().add_run("Tagesordnung").bold = True
    for i, item in enumerate(agenda_items, start=1):
        line = doc.add_paragraph(style="List Number")
        line.add_run(item.title)
        if item.description:
            sub = doc.add_paragraph(item.description)
            sub.paragraph_format.left_indent = Cm(1)

    # Schlusstext.
    closing = html_to_text(meeting.closing_text)
    if closing:
        doc.add_paragraph()
        doc.add_paragraph(closing)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
