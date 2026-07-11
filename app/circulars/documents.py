"""PDF-/Word-Erzeugung + Mail-Text für Rundschreiben.

Der Brief reitet auf dem Rechnungs-/Sitzungs-Design (Briefkopf, Adressfenster,
Farben). PDF via WeasyPrint (ImportError/OSError-Fallback beim Aufrufer), Word
via python-docx. Der Body ist Plaintext (mit ``{anrede}``/``{name}`` je
Empfänger ersetzt) — identisch für Mail und Brief.
"""
import io

from flask import render_template

from app.models import AppSetting
from app.invoices.design import get_design
from app.settings_service import (
    wg_settings, get_contact_info, get_contact_info_font_size,
    get_invoice_sender_address,
)
from app.circulars.services import render_circular_text


def current_design():
    return get_design(AppSetting.get("invoice.design", "classic"))


def _weasyprint():
    """WeasyPrint-HTML-Klasse oder None (GTK/WeasyPrint lokal nicht installiert)."""
    try:
        from weasyprint import HTML
        return HTML
    except (ImportError, OSError):
        return None


# ── Mail-Text (Plaintext) ────────────────────────────────────────────────────

def mail_body(circular, customer):
    """Plaintext-Mailtext: Body mit ersetzten Platzhaltern + WG-Signatur."""
    text = render_circular_text(circular.body, customer)
    wg = wg_settings()
    name = wg.get("name") or ""
    if name:
        text = f"{text}\n{name}"
    return text


# ── PDF (WeasyPrint-HTML) ────────────────────────────────────────────────────

def render_letter_html(circular, customer):
    """HTML eines einzelnen Rundschreiben-Briefs (für WeasyPrint), im
    Rechnungs-Stil."""
    return render_template(
        "circulars/pdf/letter.html",
        circular=circular, customer=customer,
        subject=circular.subject,
        body_text=render_circular_text(circular.body, customer),
        design=current_design(),
        contact_info=get_contact_info(),
        contact_info_font_size=get_contact_info_font_size(),
        invoice_sender_address=get_invoice_sender_address(),
        wg=wg_settings(),
    )


def render_merged_pdf(circular, recipients):
    """Gemergtes Brief-PDF (eine Seite je Empfänger) als bytes, oder None wenn
    WeasyPrint fehlt.

    Bewusst request-unabhängig (nimmt bereits geladene ``recipients``) — der
    SaaS-Async-Renderer ruft dieselbe Funktion auf. Jedes Dokument wird EINZELN
    gerendert und mit ``pypdf`` gemergt (kein WeasyPrint-``copy()`` über
    Dokumente — das crasht am gemeinsamen Logo mit ``PIL.UnidentifiedImageError``).
    """
    HTML = _weasyprint()
    if HTML is None:
        return None
    from pypdf import PdfWriter
    writer = PdfWriter()
    for rec in recipients:
        if rec.customer is None:
            continue
        pdf = HTML(string=render_letter_html(circular, rec.customer)).write_pdf()
        writer.append(io.BytesIO(pdf))
    writer.compress_identical_objects()
    buf = io.BytesIO()
    writer.write(buf)
    writer.close()
    buf.seek(0)
    return buf.read()


# ── Word (python-docx) ───────────────────────────────────────────────────────

def build_letter_docx(circular, customer):
    """Ein Rundschreiben als schlichter Brief-.docx (bytes): Briefkopf,
    Empfänger, Betreff, Fließtext."""
    from docx import Document
    from docx.shared import Pt, Cm

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2)

    wg = wg_settings()

    # Briefkopf (Kontaktinfo als Text, rechtsbündig).
    from app.schriftfuehrung.documents import html_to_text
    header = html_to_text(get_contact_info()) or (wg.get("name") or "")
    if header:
        p = doc.add_paragraph()
        p.alignment = 2  # rechts
        run = p.add_run(header)
        run.font.size = Pt(9)

    # Empfänger-Adressblock (mehrzeilig wie auf der Rechnung: Absender-Rückzeile,
    # Name, Straße, PLZ Ort, Land).
    doc.add_paragraph()
    sender_return = get_invoice_sender_address()
    if sender_return:
        sr = doc.add_paragraph()
        run = sr.add_run(sender_return)
        run.font.size = Pt(7)
    addr = doc.add_paragraph()
    addr.add_run(customer.letter_name).bold = True
    street = " ".join(p for p in (customer.strasse, customer.hausnummer) if p)
    city = " ".join(p for p in (customer.plz, customer.ort) if p)
    if street:
        addr.add_run("\n" + street)
    if city:
        addr.add_run("\n" + city)
    if customer.land and customer.land != "Österreich":
        addr.add_run("\n" + customer.land)

    # Betreff.
    doc.add_paragraph()
    subj = doc.add_paragraph()
    subj.add_run(circular.subject).bold = True

    # Fließtext (Plaintext mit ersetzten Platzhaltern) — Absätze an Leerzeilen.
    body = render_circular_text(circular.body, customer)
    for block in body.split("\n\n"):
        doc.add_paragraph(block.strip("\n"))

    if wg.get("name"):
        doc.add_paragraph(wg["name"])

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def render_merged_docx(circular, recipients):
    """Gemergtes Brief-.docx (eine Seite je Empfänger) als bytes, oder None."""
    from app.invoices.document_service import merge_docx_files
    sources = [build_letter_docx(circular, rec.customer)
               for rec in recipients if rec.customer]
    if not sources:
        return None
    return merge_docx_files(sources)
